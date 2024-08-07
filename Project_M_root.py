# Author: Thore Jördens
# Latest Change: 07.08.2024


import tkinter as tk
from tkinter import filedialog, Listbox, messagebox, ttk
from tkinterdnd2 import TkinterDnD, DND_FILES
from docx import Document
from docxcompose.composer import Composer
from PIL import Image, ImageTk
import os
import json
import threading
import time


class DocumentMergerApp:
    def __init__(self, root):
        # Hauptfenster erstellen
        self.root = root
        self.root.title("MS Word Document Merger")

        # Hintergrundbild laden
        self.background_image = Image.open(r"C:\Users\joerdens\project_M\assets\hintergrund.png")  # Pfad zum Bild anpassen
        self.background_photo = ImageTk.PhotoImage(self.background_image)

        # J&J Blue - Hintergrund
        jj_blue = "#002D72"

        # Canvas erstellen und Hintergrundbild anzeigen
        canvas = tk.Canvas(root, width=self.background_image.width, height=self.background_image.height)
        canvas.grid(row=0, column=0, sticky="nsew", columnspan=6, rowspan=6)
        canvas.create_image(0, 0, image=self.background_photo, anchor="nw")

        # Liste ausgewählter Dateien und speicherbare Gruppen erstellen
        self.file_list = []
        self.group_file = "document_groups.json"
        self.groups = self.load_all_groups()
    


        # Label für die Überschrift
        self.label = tk.Label(root, text="Drag & Drop files for merger:", font=("Sitka Heading", 16, "bold"), bg=jj_blue, fg="white")
        canvas.create_window(20, 5, anchor="nw", window=self.label)



        # X-Set für Buttons
        general_button_x = 565

        # Buttons fürs hinzufügen und organisieren
        self.add_button = tk.Button(self.root, text="Add File", command=self.add_files)
        canvas.create_window(general_button_x, 45, anchor="nw", window=self.add_button)

        self.move_up_button = tk.Button(self.root, text="Move Up", command=lambda: self.move_item(-1))
        canvas.create_window(general_button_x, 75, anchor="nw", window=self.move_up_button)

        self.move_down_button = tk.Button(self.root, text="Move Down", command=lambda: self.move_item(1))
        canvas.create_window(general_button_x, 105, anchor="nw", window=self.move_down_button)

        self.remove_button = tk.Button(self.root, text="Remove File", command=self.remove_selected)
        canvas.create_window(general_button_x, 135, anchor="nw", window=self.remove_button)

        self.merge_button = tk.Button(self.root, text="Merge Documents", command=self.merge_documents, 
                                       font=("Sitka Heading", 11, "bold"), bg="#4CAF50", fg="white", 
                                      padx=20, pady=10, width=11, height=1)
        
        canvas.create_window(general_button_x, 165, anchor="nw", window=self.merge_button)

        # Buttons fürs Speichern und laden
        self.save_group_button = tk.Button(self.root, text="Save Group", command=self.save_group)
        canvas.create_window(general_button_x, 260, anchor="nw", window=self.save_group_button)

        self.load_group_button = tk.Button(self.root, text="Load Group", command=self.load_group)
        canvas.create_window(general_button_x, 290, anchor="nw", window=self.load_group_button)


        # Dropdown-Menü für Gruppen
        self.group_selector_var = tk.StringVar(value="Select Group")
        if not self.groups:  # Überprüfen, ob das Dictionary leer ist
            self.group_selector_var.set("No Groups Available")
            options = ["No Groups Available"]
        else:
            options = list(self.groups.keys())
        
        self.group_selector = tk.OptionMenu(root, self.group_selector_var, *options)
        self.group_selector.config(width=17)  # Breite des Dropdown-Menüs einstellen
        canvas.create_window(general_button_x, 320, anchor="nw", window=self.group_selector)


        # Listbox
        self.listbox = tk.Listbox(root, selectmode=tk.SINGLE, width=85, height=24)
        canvas.create_window(20, 45, anchor="nw", window=self.listbox)

        # Drag & Drop für Listbox aktivieren
        root.drop_target_register(DND_FILES)
        root.dnd_bind('<<Drop>>', self.drop_files)


        # Überprüfung der Gruppen-Existenz in einem separaten Thread
        self.check_groups_thread = threading.Thread(target=self.check_groups_existence, daemon=True)
        self.check_groups_thread.start()


        # Coinfigure Grid
        root.grid_rowconfigure(0, weight=1)
        root.grid_columnconfigure(0, weight=1)

        self.listbox.bind("<Delete>", self.remove_selected_event)



    # Dokumente können über die Auswahl im Explorer hinzugefügt werden
    def add_files(self):
        files = filedialog.askopenfilenames(filetypes=[("Word Documents", "*.docx")])
        for file in files:
            self.file_list.append(file)
            self.update_listbox()

    # Ermöglicht Drag & Drop der Files
    def drop_files(self, event):
        files = self.root.tk.splitlist(event.data)
        for file in files:
            if file.endswith(".docx"):
                self.file_list.append(file)
        self.update_listbox()
    
    # Files aus der Liste entfernen
    def remove_selected(self):
        selected = self.listbox.curselection()
        if selected:
            self.file_list.pop(selected[0])
            self.listbox.delete(selected)

        self.update_listbox()

    def remove_selected_event(self, event):
        self.remove_selected()

    # Reihenfolge in der Listbox ändern
    def move_item(self, direction):
        selected = self.listbox.curselection()
        if selected:
            pos = selected[0]
            new_pos = pos + direction

            if 0 <= new_pos < self.listbox.size():
                # Tauschen der Elemente in der Liste
                self.file_list[pos], self.file_list[new_pos] = self.file_list[new_pos], self.file_list[pos]

                # Listbox aktualisieren
                self.update_listbox()

                # Das neue Element auswählen
                self.listbox.select_set(new_pos)

    
    # Zusammenführen der Dokumente, Benutzer wählt Speicherort
    def merge_documents(self):
        if not self.file_list:
            messagebox.showwarning("No Files", "Please add at least one Word document to merge.")
            return

        output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not output_path:
            return

        merged_document = Document()
        composer = Composer(merged_document)
        composer.append(Document(self.file_list[0]))

        for file in self.file_list[1:]:
            composer.append(Document(file))

        composer.save(output_path)
        messagebox.showinfo("Success", f"Documents merged and saved to {output_path}")

    
    # Dokuemente in listbox werden als Gruppe in einer JSON abegspeichert
    def save_group(self):
        if not self.file_list:
            messagebox.showwarning("No Files", "Please add at least one Word document to save the group.")
            return

        group_name = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON Files", "*.json")])
        if not group_name:
            return

        # Speichern der Dateipfade in der ausgewählten JSON-Datei
        with open(group_name, 'w') as group_file:
            json.dump(self.file_list, group_file)

        # Aktualisieren des internen Dictionaries
        self.groups[os.path.basename(group_name)] = self.file_list

        self.save_groups()
        self.update_group_selector()

        messagebox.showinfo("Success", f"Group saved as {group_name}")

    # Gruppe wird aus JSON geladen
    def load_group(self):
        group_name = self.group_selector_var.get()
        if group_name not in self.groups:
            messagebox.showwarning("Group Not Found", "Selected group could not be loaded.")
            return

        self.file_list = self.groups[group_name]
        self.update_listbox()

    def load_all_groups(self):
        if os.path.exists(self.group_file):
            try:
                with open(self.group_file, 'r') as group_file:
                    groups = json.load(group_file)
                    print(f"Loaded groups: {groups}")  # Debug-Ausgabe
                    return groups
            except Exception as e:
                print(f"Error loading groups: {e}")
                return {}
        else:
            print(f"Group file {self.group_file} does not exist.")  # Debug-Ausgabe
        return {}
    

    def check_groups_existence(self):
        while True:
            # Überprüfen, ob die Gruppen-Dateien noch existieren
            missing_groups = []
            for group_name, file_list in self.groups.items():
                if not os.path.exists(group_name):
                    missing_groups.append(group_name)

            # Fehlende Gruppen entfernen
            if missing_groups:
                for group_name in missing_groups:
                    del self.groups[group_name]
                # Threadsicheres Update der GUI
                self.root.after(0, self.save_groups)
                self.root.after(0, self.update_group_selector)

            # Wartezeit zwischen den Überprüfungen (z.B. alle 10 Sekunden)
            time.sleep(10)

    # Gruppen in einer JSON-Datei speichern
    def save_groups(self):
        with open(self.group_file, 'w') as group_file:
            json.dump(self.groups, group_file)

    def update_group_selector(self):
        # Aktualisieren Sie das Dropdown-Menü für die Gruppen
        self.group_selector['menu'].delete(0, 'end')
        for group_name in self.groups.keys():
            self.group_selector['menu'].add_command(label=group_name, command=lambda g=group_name: self.group_selector_var.set(g))

    def update_listbox(self):
        self.listbox.delete(0, tk.END)
        for index, file in enumerate(self.file_list, start=1):
            display_text = f"{index}: {os.path.basename(file)}"
            self.listbox.insert(tk.END, display_text)

    def load_all_groups(self):
        if os.path.exists(self.group_file):
            with open(self.group_file, 'r') as group_file:
                return json.load(group_file)
        return {}

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = DocumentMergerApp(root)
    root.mainloop()
